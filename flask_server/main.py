from flask import Flask, jsonify, request
from pymongo import MongoClient
from flask_cors import CORS
from flask_socketio import SocketIO, emit
from docx import Document
from lxml import etree
import json
import os
import hashlib
import datetime
import py7zr
import pandas as pd
import xml.etree.ElementTree as ET

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": ["http://localhost:3001", "http://localhost:3002"]}})
socketio = SocketIO(app, cors_allowed_origins=["http://localhost:3001", "http://localhost:3002"])

# MongoDB Atlas connection
client = MongoClient("mongodb+srv://ddas:ddas@sample.nnpef.mongodb.net/?retryWrites=true&w=majority&appName=sample")
db = client['Metadata']
collection = db['Metadata_collection']
users_collection = db['user_details_collection']

# File Metadata Extraction Functions
def get_file_metadata(file_path):
    try:
        stat_info = os.stat(file_path)
        return {
            "creationDate": datetime.datetime.fromtimestamp(stat_info.st_ctime).isoformat(),
            "lastModifiedDate": datetime.datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
            "filePath": file_path,
            "filePermissions": oct(stat_info.st_mode)[-3:],
            "fileOwner": stat_info.st_uid,
            "fileSize": stat_info.st_size,
            "checksumMD5": compute_checksum(file_path, 'md5'),
            "checksumSHA1": compute_checksum(file_path, 'sha1')
        }
    except Exception as e:
        return {"error": str(e)}

def compute_checksum(file_path, algo):
    hash_func = hashlib.md5() if algo == 'md5' else hashlib.sha1()
    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_func.update(chunk)
    return hash_func.hexdigest()

def get_docx_metadata(file):
    doc = Document(file)
    num_paragraphs = len(doc.paragraphs)
    return {"type": "docx", "paragraphs": num_paragraphs}

def get_xml_metadata(file):
    try:
        tree = etree.parse(file)
        root = tree.getroot()
        num_elements = len(root.findall('.//*'))
        return {"type": "xml", "elements": num_elements}
    except Exception as e:
        return {"type": "xml", "error": str(e)}

def get_json_metadata(file):
    try:
        data = json.load(file)
        num_keys = len(data) if isinstance(data, dict) else len(data)
        return {"type": "json", "keys": num_keys}
    except Exception as e:
        return {"type": "json", "error": str(e)}

def get_compression_info(file_path):
    try:
        with py7zr.SevenZipFile(file_path, 'r') as archive:
            compressed_size = sum([entry.size for entry in archive.list()])
        return {"compression": "7z", "compressedSize": compressed_size}
    except Exception as e:
        return {"compression": "unknown", "error": str(e)}

# Metadata Extraction Functions for Specific File Types
def extract_metadata_csv(file_path):
    if not os.path.exists(file_path):
        return None
    
    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path)
    df = pd.read_csv(file_path)
    num_rows, num_columns = df.shape
    
    metadata = {
        "dataset_name": file_name,
        "file_location": file_path,
        "size_bytes": file_size,
        "num_rows": num_rows,
        "num_columns": num_columns,
        "upload_date": pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S'),
        "type": "csv"
    }
    
    return metadata

def extract_metadata_docx(file_path):
    if not os.path.exists(file_path):
        return None
    
    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path)
    doc = Document(file_path)
    num_paragraphs = len(doc.paragraphs)
    
    metadata = {
        "dataset_name": file_name,
        "file_location": file_path,
        "size_bytes": file_size,
        "num_paragraphs": num_paragraphs,
        "upload_date": pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S'),
        "type": "docx"
    }
    
    return metadata

def extract_metadata_xml(file_path):
    if not os.path.exists(file_path):
        return None
    
    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path)
    tree = ET.parse(file_path)
    root = tree.getroot()
    num_elements = len(list(root.iter()))
    
    metadata = {
        "dataset_name": file_name,
        "file_location": file_path,
        "size_bytes": file_size,
        "num_elements": num_elements,
        "upload_date": pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S'),
        "type": "xml"
    }
    
    return metadata

def extract_metadata_json(file_path):
    if not os.path.exists(file_path):
        return None
    
    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path)
    with open(file_path, 'r') as f:
        data = json.load(f)
    num_elements = len(data)
    
    metadata = {
        "dataset_name": file_name,
        "file_location": file_path,
        "size_bytes": file_size,
        "num_elements": num_elements,
        "upload_date": pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S'),
        "type": "json"
    }
    
    return metadata

def upload_metadata(metadata):
    if metadata:
        try:
            existing_metadata = collection.find_one({"dataset_name": metadata['dataset_name']})
            if existing_metadata:
                print(f"Metadata for {metadata['dataset_name']} already exists.")
            else:
                collection.insert_one(metadata)
                print(f"Uploaded metadata for: {metadata['dataset_name']}")
        except Exception as e:
            print(f"Error uploading metadata: {e}")

def scan_directory(directory_path):
    datasets = {}
    for file_name in os.listdir(directory_path):
        file_path = os.path.join(directory_path, file_name)
        if file_name.endswith('.csv'):
            datasets[file_name] = (file_path, extract_metadata_csv)
        elif file_name.endswith('.docx'):
            datasets[file_name] = (file_path, extract_metadata_docx)
        elif file_name.endswith('.xml'):
            datasets[file_name] = (file_path, extract_metadata_xml)
        elif file_name.endswith('.json'):
            datasets[file_name] = (file_path, extract_metadata_json)
    return datasets

def process_directory_datasets(directory_path):
    datasets = scan_directory(directory_path)
    if not datasets:
        print("No supported files found in the directory.")
    for dataset_name, (file_path, metadata_function) in datasets.items():
        print(f"Processing dataset: {dataset_name} at {file_path}")
        metadata = metadata_function(file_path)
        upload_metadata(metadata)

# API Routes

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"message": "No file part"}), 400

    file = request.files['file']
    upload_folder = 'uploads'
    file_path = os.path.join(upload_folder, file.filename)
    
    if not os.path.exists(upload_folder):
        os.makedirs(upload_folder)
    
    file.save(file_path)
    
    # Extract metadata based on file type and manually set the type field
    if file.filename.endswith('.csv'):
        metadata = extract_metadata_csv(file_path)
        metadata_type = "csv"
    elif file.filename.endswith('.docx'):
        metadata = extract_metadata_docx(file_path)
        metadata_type = "docx"
    elif file.filename.endswith('.xml'):
        metadata = extract_metadata_xml(file_path)
        metadata_type = "xml"
    elif file.filename.endswith('.json'):
        metadata = extract_metadata_json(file_path)
        metadata_type = "json"
    else:
        return jsonify({"message": "Unsupported file type"}), 400
    
    # Adding common metadata and setting type based on file extension
    metadata.update({
        "type": metadata_type,  # Set to "csv", "xml", "json", or "docx" based on file type
        "lastModified": request.form.get('lastModified', 'Unknown'),
        **get_file_metadata(file_path)
    })
    
    # Add compression info if the file is compressed
    if file.filename.endswith('.7z'):
        metadata.update(get_compression_info(file_path))
    else:
        metadata["compression"] = "none"
    
    collection.insert_one(metadata)
    return jsonify({"message": "File successfully uploaded"}), 200

@app.route('/api/metadata', methods=['GET'])
def get_metadata():
    format_filter = request.args.get('format', 'all')
    query = {} if format_filter == 'all' else {'type': format_filter}
    metadata_list = list(collection.find(query, {'_id': 0}))
    return jsonify(metadata_list)

@app.route('/api/metadata/<string:filename>', methods=['GET'])
def get_metadata_by_filename(filename):
    metadata = collection.find_one({"dataset_name": filename}, {'_id': 0})
    if metadata:
        return jsonify(metadata)
    else:
        return jsonify({"error": "Metadata not found"}), 404

@app.route('/api/process_directory', methods=['POST'])
def process_directory():
    directory_path = request.json.get('directory_path')
    if directory_path:
        process_directory_datasets(directory_path)
        return jsonify({"message": "Directory processed successfully"})
    else:
        return jsonify({"error": "Directory path is required"}), 400

@app.route('/api/signup', methods=['POST'])
def signup():
    data = request.json
    if users_collection.find_one({"email": data["email"]}):
        return jsonify({"message": "User already exists"}), 409
    users_collection.insert_one(data)
    return jsonify({"message": "User registered successfully"})

@app.route('/api/login', methods=['POST'])
def login():
    data = request.json
    username = data.get('username')
    password = data.get('password')
    user = users_collection.find_one({"username": username})
    if user and user["password"] == password:
        user["_id"] = str(user["_id"])
        role = user.get("role", "").strip().lower()
        return jsonify({"message": "Login successful", "role": role, "user_id": user["_id"]})
    else:
        return jsonify({"message": "Invalid credentials", "success": False}), 401

@socketio.on('connect')
def handle_connect():
    print("Client connected")
    emit('response', {'data': 'Connected to WebSocket server'})

@socketio.on('message')
def handle_message(data):
    print('Message received:', data)
    emit('response', {'data': 'Message received'})

if __name__ == '__main__':
    socketio.run(app, host='0.0.0.0', port=5000, debug=True)
